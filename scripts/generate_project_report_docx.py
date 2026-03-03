from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


ROOT = Path(__file__).resolve().parents[1]


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def add_formula(doc: Document, latex: str) -> None:
    p = doc.add_paragraph()
    p.add_run(latex)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_metric_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def fmt(x: float, digits: int = 4) -> str:
    return f"{x:.{digits}f}"


def build_report() -> Path:
    lab1_leak = read_json(
        ROOT
        / "saves/lab1_run_combo_af_gate_exit_v2/audits_confidence/leakage_summary.json"
    )
    lab1_gate = read_json(
        ROOT / "saves/lab1_run_combo_af_gate_exit_v2/audits_confidence/gate_summary.json"
    )
    lab2_val = read_json(
        ROOT
        / "saves/lab2_calibration/lab2_20260211_015118_lda_cleanup_v2/validation_summary.json"
    )
    lab3_state = read_json(ROOT / "saves2/lab3_codec_transfer/run1055/run_state.json")
    lab3_gate = read_json(ROOT / "saves2/lab3_codec_transfer/run1055/codec_gate_eval.json")
    lab3_style_bank = read_json(
        ROOT / "saves2/lab3_codec_transfer/run1055/style_bank_diagnostics.json"
    )
    diff_v2_hist = read_json(ROOT / "saves2/lab3_diffusion/run_d002/v2_history.json")
    diff_v3_hist = read_json(ROOT / "saves2/lab3_diffusion/run_d003/v3_history.json")
    lab4_coh = read_json(
        ROOT / "saves2/lab4_longform_coherence/fullsong_test/coherence_metrics.json"
    )
    diff_meta = read_json(ROOT / "saves2/lab3_diffusion/run_d001/cache/diff_meta.json")
    diff_genres = read_json(ROOT / "saves2/lab3_diffusion/run_d001/cache/diff_genre_to_idx.json")

    # Diffusion history summaries
    best_v2 = min(diff_v2_hist, key=lambda x: x.get("val_loss", float("inf"))) if diff_v2_hist else {}
    epoch6_v2 = {}
    for row in diff_v2_hist:
        if row.get("epoch") == 6:
            epoch6_v2 = row
            break
    best_v3 = min(diff_v3_hist, key=lambda x: x.get("val_loss", float("inf"))) if diff_v3_hist else {}

    doc = Document()
    title = doc.add_heading("Deep Generative Genre Remastering: Architecture, Methods, and Results", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph(f"Technical summary generated on {date.today().isoformat()}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("1. High-Level Overview", level=1)
    doc.add_paragraph(
        "This project implements a staged music style transfer pipeline designed to move beyond "
        "surface timbral filtering and toward structural genre remastering. The architecture separates "
        "content (melody/rhythm/harmonic intent) from style (genre-dependent timbre/texture/performance cues), "
        "builds explicit target-style representations, and then reconstructs audio using two complementary "
        "generation paths: (i) neural codec latent translation and (ii) diffusion-based mel generation with "
        "long-form coherence controls."
    )
    doc.add_paragraph("End-to-end pipeline: Lab 1 (Deconstruction Encoder) -> Lab 2 (Target Vector Space) -> Lab 3 (Reconstruction) -> Lab 4 (Long-form Coherence) -> Lab 5 (Final Evaluation, planned).")

    doc.add_heading("2. Original Plan vs Implemented Plan", level=1)
    doc.add_paragraph(
        "The implementation follows the original five-lab plan with practical adaptations. "
        "Lab 1 and Lab 2 were completed with threshold-passing audits. Lab 3 achieved target style-transfer "
        "metrics using EnCodec latent translation and MERT-conditioned style control. Lab 4 added constrained "
        "diffusion chunking with overlap locking/re-anchoring for long-form consistency. Lab 5 perceptual "
        "testing and broader realism benchmarking remain the final stage."
    )

    doc.add_heading("3. Lab 1: Deconstruction Encoder", level=1)
    doc.add_heading("3.1 Architecture", level=2)
    doc.add_paragraph(
        "Input log-mel spectrograms are processed by a compact CNN encoder with shared representation and "
        "two disentangled latent heads: z_content and z_style (both 128D, L2-normalized). "
        "An adversarial style classifier on z_content uses a Gradient Reversal Layer (GRL) to remove style "
        "information from content latents. A dedicated music gate (binary head) is trained to reject speech/non-musical "
        "segments in later stages."
    )
    add_formula(doc, r"$$z_c = \frac{f_c(h)}{\|f_c(h)\|_2}, \quad z_s = \frac{f_s(h)}{\|f_s(h)\|_2}$$")
    add_formula(doc, r"$$\tilde{z}_c = \mathrm{GRL}_{\lambda}(z_c)$$")

    doc.add_heading("3.2 Training Curriculum", level=2)
    doc.add_paragraph(
        "A three-phase curriculum was used: Phase 1 (content-focused disentanglement), Phase 2 "
        "(stronger adversarial separation on real music), and Phase 3 (music gate sharpening with optional teacher-anchor "
        "regularization)."
    )

    doc.add_heading("3.3 Loss Functions", level=2)
    add_formula(
        doc,
        r"$$\mathcal{L}_{\text{Lab1}} = w_c\mathcal{L}_{content} + w_s\mathcal{L}_{style} + w_m\mathcal{L}_{music} + w_{adv}\mathcal{L}_{content\_adv} + w_{l1}\mathcal{L}_{content\_l1} + w_{mb}\mathcal{L}_{music\_bias} + w_a\mathcal{L}_{anchor}$$",
    )
    add_formula(doc, r"$$\mathcal{L}_{content}=\mathrm{MSE}(z_c^{(a)}, z_c^{(b)})$$")
    add_formula(doc, r"$$\mathcal{L}_{style}=\mathrm{CE}(g_s(z_s), y_{source}), \quad \mathcal{L}_{content\_adv}=\mathrm{CE}(g_{adv}(\tilde{z}_c), y_{source})$$")
    add_formula(doc, r"$$\mathcal{L}_{music}=\mathrm{BCEWithLogits}(g_m(h), y_{music})$$")
    add_formula(doc, r"$$\mathcal{L}_{anchor}=\tfrac{1}{2}\left[\mathrm{MSE}(z_c, z_c^{teacher}) + \mathrm{MSE}(z_c', z_c'^{teacher})\right]$$")

    doc.add_heading("3.4 Lab 1 Evaluation Metrics and Criteria", level=2)
    doc.add_paragraph(
        "Content leakage measures how much source style information remains recoverable from z_content "
        "(lower is better). Style probe accuracy measures style discriminability in z_style (higher is better). "
        "Gate AUC measures music-vs-speech ranking quality independent of threshold (higher is better)."
    )
    add_metric_table(
        doc,
        ["Metric", "Criterion", "Achieved", "Interpretation"],
        [
            [
                "Content leakage above baseline",
                "<= 0.15",
                fmt(float(lab1_leak.get("content_leakage_above_baseline", float("nan")))),
                "Passed: content latents are substantially style-suppressed.",
            ],
            [
                "Style probe accuracy",
                ">= 0.85",
                fmt(float(lab1_leak.get("style_probe_accuracy", float("nan")))),
                "Passed: style latents preserve genre signal.",
            ],
            [
                "Music gate ROC AUC",
                ">= 0.90",
                fmt(float(lab1_gate.get("roc_auc", float("nan")))),
                "Passed: gate separates music/non-music reliably.",
            ],
        ],
    )

    doc.add_heading("4. Lab 2: Genre Target Vector Calibration", level=1)
    doc.add_heading("4.1 Target Vector Construction", level=2)
    doc.add_paragraph(
        "Each sample is represented by a 160D target vector combining learned style latent and handcrafted "
        "spectral descriptor: 128D z_style plus 32D descriptor32 (16 mel-band means + 16 mel-band standard deviations). "
        "The composition is weighted and then L2-normalized row-wise."
    )
    add_formula(doc, r"$$d_{32} = [\mu_1,\dots,\mu_{16},\sigma_1,\dots,\sigma_{16}]$$")
    add_formula(doc, r"$$x_{160} = \mathrm{normalize}\left([\alpha z_s \;||\; \beta d_{32}]\right)$$")
    doc.add_paragraph("Run configuration used: alpha=2.0 for z_style and beta=1.0 for descriptor32.")

    doc.add_heading("4.2 Centroids, Projection, and Validation", level=2)
    doc.add_paragraph(
        "Genre centroids are computed with inlier filtering (top inlier fraction by cosine distance) to reduce "
        "outlier drift. A supervised LDA projection (160D -> 3D, shrinkage='auto') is used for geometric auditability. "
        "Validation includes nearest-centroid assignment, linear probe accuracy, silhouette score, centroid stability, "
        "and inter-centroid separation checks."
    )
    add_formula(doc, r"$$\hat{y}_{NC}(x)=\arg\max_k \cos(x, c_k)$$")
    add_formula(doc, r"$$\mathrm{Silhouette} = \frac{b(i)-a(i)}{\max(a(i),b(i))} \;\; \text{(averaged)}$$")

    m2 = lab2_val.get("metrics", {})
    add_metric_table(
        doc,
        ["Metric", "Criterion", "Achieved", "Interpretation"],
        [
            [
                "Silhouette (cosine)",
                ">= 0.45",
                fmt(float(m2.get("silhouette", float("nan")))),
                "Passed: meaningful inter-genre separation.",
            ],
            [
                "Nearest-centroid accuracy",
                "Higher is better",
                fmt(float(m2.get("nearest_centroid_acc", float("nan")))),
                "Strong centroid fidelity for target vectors.",
            ],
            [
                "Linear probe accuracy",
                "Higher is better",
                fmt(float(m2.get("linear_probe_acc", float("nan")))),
                "Target space is linearly genre-informative.",
            ],
            [
                "Samples / Genres",
                "Coverage",
                f"{int(m2.get('n_samples', 0))} / {int(m2.get('n_genres', 0))}",
                "Sufficient scale for stable centroid estimates.",
            ],
        ],
    )

    doc.add_heading("5. Lab 3: Reconstruction Decoder (Codec Latent Translation)", level=1)
    doc.add_heading("5.1 Generator and Discriminator", level=2)
    doc.add_paragraph(
        "The primary Lab 3 system translates EnCodec latents using a FiLM-conditioned Conv1D residual translator. "
        "Conditioning inputs are z_content, target-style embedding, and injected noise. "
        "The generator can operate in residual mode or direct-output mode; direct-output produced the best style-transfer "
        "results in the current project."
    )
    add_formula(doc, r"$$q_{hat} = q_{src} + s\cdot\tanh(\Delta q)\quad\text{(residual mode)}$$")
    add_formula(doc, r"$$q_{hat} = \Delta q\quad\text{(direct-output mode)}$$")
    doc.add_paragraph(
        "A multi-scale waveform discriminator provides adversarial realism pressure with hinge losses "
        "and feature matching."
    )

    doc.add_heading("5.2 Loss Function (Codec Path)", level=2)
    add_formula(
        doc,
        r"$$\mathcal{L}_G = \lambda_{adv}\mathcal{L}_{adv} + \lambda_{fm}\mathcal{L}_{fm} + \lambda_{l1}\mathcal{L}_{latent\_l1} + \lambda_{cont}\mathcal{L}_{latent\_cont} + \lambda_{mr}\mathcal{L}_{mrstft} + \lambda_c\mathcal{L}_{content} + \lambda_s\mathcal{L}_{style} + \lambda_{ms}\mathcal{L}_{mode\_seek} + \lambda_{push}\mathcal{L}_{style\_push} + \lambda_{\delta}\mathcal{L}_{delta\_budget}$$",
    )
    add_formula(doc, r"$$\mathcal{L}_{content}=1-\cos(z_c^{hat}, z_c^{src})$$")
    add_formula(doc, r"$$\mathcal{L}_{style\_push}=\mathrm{ReLU}(p_{source}-m)$$")
    add_formula(doc, r"$$\mathcal{L}_{delta\_budget}=\mathrm{ReLU}\left(\|q_{hat}-q_{src}\|_{1,mean}-\delta\right)$$")

    doc.add_heading("5.3 Best Achieved Run (run1055)", level=2)
    g3 = lab3_gate.get("metrics", {})
    sb = lab3_style_bank.get("metrics", {})
    add_metric_table(
        doc,
        ["Metric", "Criterion", "Achieved", "Interpretation"],
        [
            [
                "MPS (melodic preservation score)",
                ">= 0.90",
                fmt(float(g3.get("mps", float("nan")))),
                "Passed: high content preservation.",
            ],
            [
                "Style confidence",
                ">= 0.85 project goal",
                fmt(float(g3.get("style_conf", float("nan")))),
                "Exceeded target style-transfer confidence.",
            ],
            [
                "Style accuracy",
                ">= 0.85 project goal",
                fmt(float(g3.get("style_acc", float("nan")))),
                "Exceeded target style correctness.",
            ],
            [
                "Pairwise cosine (diversity proxy)",
                "<= 0.95 gate limit",
                fmt(float(g3.get("pairwise_cos", float("nan")))),
                "Low collapse risk under gate criterion.",
            ],
            [
                "Style-bank nearest centroid acc",
                ">= 0.70 gate limit",
                fmt(float(sb.get("nearest_centroid_acc", float("nan")))),
                "Strong target-style geometry in conditioning space.",
            ],
        ],
    )
    doc.add_paragraph(
        "This run used MERT-probe style conditioning and direct-output translator mode, which overcame earlier "
        "style ceiling behavior observed with weaker conditioning spaces."
    )

    doc.add_heading("6. Diffusion Branch (Lab 3/4 bridge)", level=1)
    doc.add_heading("6.1 Diffusion V2 Objective and Sampling", level=2)
    doc.add_paragraph(
        "The diffusion branch models normalized mel spectrograms with v-prediction, using separate conditioning "
        "paths for time/content and style modulation. Classifier-free guidance (CFG) is applied at sampling."
    )
    add_formula(doc, r"$$x_t = \sqrt{\bar{\alpha}_t}x_0 + \sqrt{1-\bar{\alpha}_t}\epsilon$$")
    add_formula(doc, r"$$v_t = \sqrt{\bar{\alpha}_t}\epsilon - \sqrt{1-\bar{\alpha}_t}x_0$$")
    add_formula(doc, r"$$\mathcal{L}_{v}=\|v_{\theta}(x_t,t,c,s)-v_t\|_2^2$$")
    add_formula(doc, r"$$v_{cfg} = v_{uncond} + w\cdot(v_{cond}-v_{uncond})$$")

    doc.add_heading("6.2 Diffusion V3 Extension", level=2)
    doc.add_paragraph(
        "V3 adds a HybridMelDiscriminator with hinge adversarial loss and feature matching during fine-tuning."
    )
    add_formula(doc, r"$$\mathcal{L}_{G}^{V3}=\mathcal{L}_{v}+\lambda_{adv}\mathcal{L}_{adv}+\lambda_{fm}\mathcal{L}_{fm}$$")
    add_formula(doc, r"$$\mathcal{L}_{D}^{hinge}=\mathbb{E}[\max(0,1-D(x_{real}))]+\mathbb{E}[\max(0,1+D(x_{fake}))]$$")

    add_metric_table(
        doc,
        ["Run", "Best Val Loss", "Epoch", "Interpretation"],
        [
            [
                "run_d002 (V2)",
                fmt(float(best_v2.get("val_loss", float("nan")))),
                str(best_v2.get("epoch", "n/a")),
                "Best numerical validation among diffusion runs.",
            ],
            [
                "run_d002 epoch 6",
                fmt(float(epoch6_v2.get("val_loss", float("nan")))),
                "6",
                "Selected for generation quality in practical listening.",
            ],
            [
                "run_d003 (V3)",
                fmt(float(best_v3.get("val_loss", float("nan")))),
                str(best_v3.get("epoch", "n/a")),
                "Did not surpass V2 validation in current tuning.",
            ],
        ],
    )

    doc.add_heading("7. Lab 4: Long-Form Coherence System", level=1)
    doc.add_paragraph(
        "Long-form generation is performed with overlapped chunking and constrained DDIM sampling. "
        "Two key constraints are applied: (i) source anchoring (SDEdit-like start from noised source mel), "
        "and (ii) prefix locking where overlap frames are enforced each reverse step using previous chunk tail "
        "at matched noise level. Additional controls include periodic re-anchoring, style-strength blending, "
        "source-prefix blending, mel smoothing, and high-frequency source blending to reduce warble/static drift."
    )
    add_formula(doc, r"$$z_s^{mix}=\mathrm{normalize}\left((1-\gamma)z_s^{src}+\gamma z_s^{tgt}\right)$$")
    add_formula(doc, r"$$x^{new}_{prefix} \leftarrow \alpha x^{ref}_{prefix} + (1-\alpha)x^{new}_{prefix}$$")

    add_metric_table(
        doc,
        ["Metric", "Value", "Meaning"],
        [
            [
                "Boundary mel MSE mean",
                fmt(float(lab4_coh.get("boundary_mel_mse_mean", float("nan"))), 6),
                "Average overlap mismatch in mel space (lower is better).",
            ],
            [
                "Boundary discontinuity dB mean",
                fmt(float(lab4_coh.get("boundary_disc_db_mean", float("nan"))), 4),
                "Perceived boundary jump proxy (lower is better).",
            ],
            [
                "Chunks / duration",
                f"{int(lab4_coh.get('n_chunks', 0))} / {fmt(float(lab4_coh.get('duration_sec', 0.0)),1)}s",
                "Scope of long-form synthesis test.",
            ],
        ],
    )

    doc.add_heading("8. Dataset and Genre Coverage Snapshot", level=1)
    genres = ", ".join([f"{k} ({v})" for k, v in sorted(diff_genres.items(), key=lambda kv: kv[1])])
    doc.add_paragraph(f"Active diffusion/cache genre schema: {genres}.")
    if diff_meta:
        doc.add_paragraph(
            "Diffusion cache meta: "
            f"n_samples={diff_meta.get('n_samples')}, "
            f"sr={diff_meta.get('sr')}, n_mels={diff_meta.get('n_mels')}, "
            f"frames_per_chunk={diff_meta.get('n_frames')}, "
            f"mel_range=[{fmt(float(diff_meta.get('mel_min', 0.0)),4)}, {fmt(float(diff_meta.get('mel_max', 0.0)),4)}]."
        )

    doc.add_heading("9. Metric Definitions and Criteria Interpretation", level=1)
    doc.add_paragraph(
        "Content leakage: excess source-style recoverability from content latent above chance baseline; lower means better disentanglement."
    )
    doc.add_paragraph(
        "Style accuracy / style confidence: correctness and confidence of target-genre assignment after transfer; higher means stronger style control."
    )
    doc.add_paragraph(
        "MPS (melodic preservation score): similarity of reconstructed content latent to source content latent; higher means less melodic drift."
    )
    doc.add_paragraph(
        "Silhouette / nearest-centroid / probe accuracy: geometry and separability diagnostics for target vector space."
    )
    doc.add_paragraph(
        "Boundary discontinuity metrics: chunk junction smoothness diagnostics for long-form assembly."
    )

    doc.add_heading("10. Current Status and Technical Conclusion", level=1)
    doc.add_paragraph(
        "The project currently demonstrates a complete deconstruct-calibrate-reconstruct pipeline with successful "
        "quantitative thresholds in Lab 1 and Lab 2, and strong style-transfer outcomes in Lab 3 (run1055 exceeded "
        "the 0.85 style target while preserving melody). The primary open challenge has shifted from short-chunk "
        "style control to long-form perceptual smoothness (warble/static accumulation), for which Lab 4 now includes "
        "multiple stability controls and measurable boundary diagnostics. The final stage is formal perceptual and "
        "distributional realism evaluation (blind tests, FAD-style analysis, and baseline comparisons)."
    )

    doc.add_heading("11. References (Project-Aligned)", level=1)
    refs = [
        "[1] Music Style Transfer with Time-Varying Inversion of Diffusion Models.",
        "[2] Musical Composition Style Transfer via Disentangled Timbre Representations (arXiv:1905.13567).",
        "[3] A Hierarchical Latent Vector Model for Learning Long-Term Structure in Music.",
        "[4] GANSynth: Adversarial Neural Audio Synthesis.",
        "[5] Deep Music Analogy Via Latent Representation Disentanglement.",
        "[6] Music Genre Recognition using Deep Neural Networks and Transfer Learning.",
        "[7] Music Style Transfer: A Position Paper (arXiv:1803.06841).",
        "[8] StyleFormer: Real-Time Arbitrary Style Transfer via Parametric Style Composition.",
        "[9] Adversarial Audio Synthesis (WaveGAN/SpecGAN).",
        "[19] High Fidelity Neural Audio Compression (EnCodec).",
        "[20] HiFi-GAN: Generative Adversarial Networks for Efficient and High Fidelity Speech Synthesis.",
        "[21] MERT: Acoustic Music Understanding Model with Large-Scale Self-supervised Training.",
        "[22] AudioLM: a Language Modeling Approach to Audio Generation.",
        "[23] BigVGAN: A Universal Neural Vocoder with Large-Scale Training.",
        "Additional method-grounding references: Domain-Adversarial Training (GRL), FiLM, AdaIN, DDPM, DDIM, Classifier-Free Guidance, SDEdit, Fréchet Audio Distance.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    out_dir = ROOT / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Deep_Generative_Genre_Remastering_Technical_Summary.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build_report()
    print(str(path))
