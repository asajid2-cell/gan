from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
TEX_OUT = ROOT / "dggr_speaker_script.tex"
DOCX_OUT = ROOT / "dggr_speaker_script.docx"


EQUATION_EXPLANATIONS = {
    6: [
        "Equation 1 states the overall task: start from a source audio example x_src and produce a remastered output x_hat_tgt that should belong to a different target genre.",
        "Equation 2 expresses the two constraints that make the task hard. The content term means the generated output should preserve the recognizable musical identity of the source, such as melody, phrase contour, and rhythmic backbone.",
        "The manifold term means the output should lie inside the set of samples that sound plausible for the target genre. Here M denotes the target-genre manifold, and tgt denotes the requested target style.",
    ],
    10: [
        "The first equation says that the encoder E maps an input audio example x into two separate latent codes: z_c for content and z_s for style.",
        "In this slide, z_c is intended to store musically stable information such as melody, rhythm, and structural identity, while z_s is intended to store genre-bearing information such as timbre, articulation, texture, and arrangement cues.",
        "The loss equation defines the training objective as a weighted sum of three parts. L_content rewards content preservation, L_style rewards style informativeness, and L_adv is the adversarial term that pushes the model to hide style information inside the content code.",
        "The coefficients lambda_c, lambda_s, and lambda_a are scalar weights that control the influence of each objective during training.",
    ],
    11: [
        "The first equation defines our target style vector. The notation [a || b] means concatenation: we join a learned style embedding z_s with an auxiliary 32-dimensional descriptor d_32 and then normalize the result.",
        "The factor 2.0 in front of z_s and the factor 1.0 in front of d_32 indicate relative weighting before normalization. In practice, this means the learned style latent is emphasized more strongly than the hand-crafted descriptor block.",
        "The second equation defines a genre centroid mu_g. For a genre g, we average the vectors v(x) from all samples x in the set S_g belonging to that genre.",
        "Here |S_g| is the number of samples in genre g, and v(x) is the final style representation extracted for sample x. The centroid acts as a stable geometric summary of the target genre.",
    ],
    12: [
        "The FiLM equation describes feature-wise linear modulation. A feature tensor F is scaled by gamma(c) and shifted by beta(c), where both modulation terms are learned from a condition vector c.",
        "Here c can represent conditioning information such as content, time, or another control signal. The symbol for element-wise multiplication means each feature channel can be modulated independently.",
        "The AdaIN equation applies a similar idea to normalized features. IN(F) denotes instance normalization, and the condition s produces learned scale and shift parameters through gamma(s) and beta(s).",
        "In practice, FiLM is useful when the condition should guide internal structure throughout the network, while AdaIN is often useful when the condition behaves more like style or timbral texture.",
    ],
    14: [
        "The forward-process equation defines how a clean sample x_0 is corrupted into a noisy sample x_t at timestep t. The term alpha_bar_t controls how much of the original signal remains, and the remaining factor sets the noise covariance.",
        "The notation q(x_t | x_0) means the probability distribution of the noisy sample conditioned on the clean sample. The Gaussian form makes it possible to sample noisy versions of the input at arbitrary timesteps.",
        "The guidance equation defines classifier-free guidance. epsilon_theta(x_t, t, null) is the model's unconditional noise prediction, while epsilon_theta(x_t, t, c) is the conditional prediction under condition c.",
        "The scalar w is the guidance scale. Increasing w pushes the reverse process more strongly toward the condition, which usually strengthens the edit but can also make artifacts or content loss more likely.",
    ],
    15: [
        "The first equation describes source anchoring for chunk k. Instead of starting from pure noise, the system begins from a noised version of the original source chunk x_0^(k).",
        "Here alpha(t) controls how much of the original source chunk remains, sigma(t) controls the noise level, and epsilon is random Gaussian noise. This preserves large-scale structure while still allowing stylistic editing.",
        "The second equation describes prefix locking across chunks. The prefix of the current chunk x_t^(k) is forced to match the tail of the previous chunk x_t^(k-1) during sampling.",
        "The prefix operator refers to the overlapping start of the current chunk, and the tail operator refers to the overlapping end of the previous chunk. This is meant to reduce boundary discontinuities and seam artifacts.",
    ],
    26: [
        "This equation defines the full codec-branch training objective as a weighted sum of realism losses, reconstruction losses, preservation losses, and anti-identity pressure.",
        "The term L_GAN is the adversarial loss that encourages the output to sound realistic to the discriminator. L_feat is the feature-matching loss, which stabilizes adversarial training by matching internal discriminator activations.",
        "The term |q_hat minus q_star|_1 is an L1 distance between the predicted codec representation q_hat and the target codec representation q_star. L_MRSTFT is the multi-resolution STFT loss, which helps preserve spectral structure across multiple time-frequency scales.",
        "The terms L_content and L_style enforce the remastering objective directly: preserve the source musical identity while moving toward the target genre. L_push discourages trivial identity solutions that stay too close to the source.",
        "Each lambda coefficient controls the strength of one component. The final behavior depends on balancing these weights rather than maximizing any single term in isolation.",
    ],
}


SLIDES = [
    {
        "num": 1,
        "title": "Music Style Transfer and Deep Generative Genre Remastering",
        "time": "1-2 min",
        "goal": "Open the talk, set expectations, and make it clear that the presentation is divided into a general lecture and a project case study.",
        "script": [
            "I would open by telling the audience that the talk has two distinct jobs. The first job is educational: we want to explain what music style transfer is, why it is technically hard, and what model families the literature uses to approach it. The second job is evaluative: once we have given that general framework, we want to use our own project, Deep Generative Genre Remastering or DGGR, as a concrete case study that either succeeds or fails along those same dimensions.",
            "I would also say, right away, that this is not a talk about adding a surface filter to a song. The deeper question is whether a system can make a piece of music sound as if it were originally conceived in a different genre while still preserving the identity of the source. That distinction sets up everything that follows.",
        ],
        "transition": "With that framing in place, the next slide explains how the talk is organized and why we split the content this way.",
    },
    {
        "num": 2,
        "title": "Talk Map",
        "time": "1 min",
        "goal": "Walk the audience through the flow of the lecture so the later project details feel motivated rather than abrupt.",
        "script": [
            "On this slide I would tell the audience that the lecture half comes first on purpose. If we start directly with our architecture, then every design choice looks arbitrary. By starting with the field, we can motivate why disentanglement matters, why target spaces matter, why codec editing and diffusion are both reasonable, and why long-form coherence becomes a systems problem.",
            "Then I would preview the case-study half. Once the audience understands the general tools, we can ask whether our own pipeline is a good answer to the problem. That way the project section does not feel like a disconnected engineering demo; it feels like a concrete instantiation of the general theory.",
        ],
        "transition": "The first technical question, then, is why music style transfer is genuinely hard in the first place.",
    },
    {
        "num": 3,
        "title": "Part I: General lecture",
        "time": "0.5 min",
        "goal": "Mark the shift into lecture mode and briefly reset the audience.",
        "script": [
            "At this section divider I would pause and explicitly say that the next run of slides is about the field itself, not yet our implementation. That helps the audience mentally switch from project mode into concept mode.",
        ],
        "transition": "The natural starting point is the difficulty of the task itself.",
    },
    {
        "num": 4,
        "title": "Why Music Style Transfer Is Hard",
        "time": "2 min",
        "goal": "Explain why naive audio style transfer fails and define the main technical subproblems.",
        "script": [
            "Here I would start with the central critique from the literature: a lot of early work changed surface timbre without changing the underlying musical organization. That is the coat-of-paint problem. A piano passage pushed through a violin-like texture is not the same thing as a phrase that sounds arranged for strings.",
            "Then I would emphasize that music is multilevel. Melody, harmony, rhythm, instrumentation, articulation, and production cues are all mixed into the same audio stream. If a model edits one of these badly, it often damages another. This is why style transfer in music is structurally different from a simpler image-style analogy.",
            "Finally, I would list the four hard subproblems: separate content from style, represent the target genre in a usable latent space, reconstruct realistic audio, and maintain consistency over time. Those become the backbone of the rest of the lecture and also the backbone of our project.",
        ],
        "transition": "To solve those problems rigorously, we need precise definitions of content, style, and genre.",
    },
    {
        "num": 5,
        "title": "Content, Style, and Genre",
        "time": "2 min",
        "goal": "Give working definitions that will support the later architecture discussion.",
        "script": [
            "For this slide I would define content as the parts of the song that listeners still recognize after an instrumentation change: the melody, the harmonic movement, the rhythmic backbone, and the phrase structure. These are the elements we want to preserve.",
            "I would define style as the way those ideas are performed and produced: timbre, articulation, groove feel, attack and decay behavior, arrangement density, and production texture. These are the things we want to change when we remaster into a new genre.",
            "Genre is broader. It is not just a bundle of acoustic features. It is a cultural and statistical manifold shaped by recurrent stylistic patterns, conventions, and listener expectations. That is why genre is harder than timbre, but still learnable enough for machine learning systems to exploit.",
        ],
        "transition": "Once we have those definitions, we can state the problem more formally.",
    },
    {
        "num": 6,
        "title": "Problem Formulation",
        "time": "2 min",
        "goal": "Present the core objective mathematically and explain what the equations mean in plain language.",
        "equations": [
            r"x_{\mathrm{src}} \rightarrow \hat{x}_{\mathrm{tgt}}",
            r"\mathrm{content}(\hat{x}_{\mathrm{tgt}}) \approx \mathrm{content}(x_{\mathrm{src}}), \qquad \hat{x}_{\mathrm{tgt}} \in \mathcal{M}_{\mathrm{genre=tgt}}",
        ],
        "script": [
            "I would read the first expression very simply: we start from a source sample and want to generate a target-domain sample. But the second line is where the real difficulty appears. We are imposing two constraints at once. First, the content of the generated sample should remain close to the content of the source. Second, the generated sample should lie on the manifold of the target genre.",
            "This is the key tension in the whole talk. If we push too hard toward the target genre, we lose the melody or rhythmic identity. If we preserve the source too strongly, we only get a weak stylistic edit. Every model family we discuss can be understood as a different way of managing that tradeoff.",
        ],
        "transition": "The next question is which representation gives us the best leverage over those competing objectives.",
    },
    {
        "num": 7,
        "title": "Representation Levels in Music ML",
        "time": "2 min",
        "goal": "Explain why representation choice is foundational and connect it to later design decisions.",
        "script": [
            "I would explain that waveform-level modeling gives maximal acoustic fidelity, but it is also the hardest space in which to perform semantic edits. A raw one-dimensional signal does not expose harmony, onset structure, or spectral texture in a way that is easy for the model to manipulate conditionally.",
            "Time-frequency representations such as log-mel spectrograms, chroma, onset maps, and beat grids expose more structure. They are not perfect, but they make it easier to talk about rhythm, harmonic centroids, and timbral energy. Symbolic representations like MIDI are even better for long-range structure, but they usually lose the realism that matters in a remastering task.",
            "The lesson is that strong systems often combine levels. They preserve or analyze structure in a spectrogram or symbolic-like space, then restore waveform realism through a decoder, vocoder, or neural codec. That exact pattern appears later in DGGR.",
        ],
        "transition": "Before building anything, though, we need to understand how the field got here historically.",
    },
    {
        "num": 8,
        "title": "Historical Evolution of the Field",
        "time": "2 min",
        "goal": "Place the project in a research timeline and show why staged systems became common.",
        "script": [
            "On this slide I would summarize the field in phases. Early work relied on direct mappings, signal transforms, or image-style analogies applied to spectrograms. These approaches were often brittle and lacked semantic control.",
            "The next phase focused on disentanglement. Researchers began trying to separate pitch, rhythm, timbre, and other factors in latent space. That shift was important because it recognized that style transfer is really a representation-learning problem as much as a synthesis problem.",
            "More recent work added stronger generators such as GANs, codecs, universal vocoders, and diffusion models. But once realism improved, long-form consistency became the harder challenge. That is why modern systems increasingly look like staged pipelines rather than single end-to-end autoencoders.",
        ],
        "transition": "The first major tool in that staged pipeline view is disentanglement.",
    },
    {
        "num": 9,
        "title": "What Evaluation Must Measure",
        "time": "2 min",
        "goal": "Establish evaluation criteria early so the audience knows how later results should be interpreted.",
        "script": [
            "I would use this slide to make a simple but important point: there is no single score that tells us whether style transfer worked. We need at least four axes. First, content preservation: did the song identity survive? Second, style fidelity: did the output really move toward the target genre? Third, audio realism: does the output sound plausible as audio rather than as a synthetic artifact? And fourth, for longer sequences, coherence: does the result stay stable over time?",
            "This matters because many misleading claims in generative audio happen when one of these dimensions is ignored. For example, a model can get high target-genre confidence by generating noisy or stereotyped outputs. Or it can preserve the melody well by barely changing the style at all. So later, when we report our own metrics, we want the audience to already understand that each number only answers one part of the overall question.",
        ],
        "transition": "With the evaluation framework in place, we can examine the first core tool: disentanglement.",
    },
    {
        "num": 10,
        "title": "Tool 1: Disentanglement",
        "time": "2-3 min",
        "goal": "Explain the core disentanglement objective and the role of the gradient reversal layer.",
        "equations": [
            r"E(x)\rightarrow (z_c,z_s)",
            r"\mathcal{L}=\lambda_c \mathcal{L}_{content}+\lambda_s \mathcal{L}_{style}+\lambda_a \mathcal{L}_{adv}",
        ],
        "script": [
            "I would explain the notation first. The encoder maps an input sample into two latent codes: a content code z sub c and a style code z sub s. The loss function then tries to shape those codes into the behaviors we want. The content term rewards invariance, the style term rewards style informativeness, and the adversarial term punishes the content code when style is recoverable from it.",
            "The gradient reversal layer is the critical mechanism here. On the forward pass, it does nothing. But on the backward pass, it flips the sign of the gradient coming from the adversarial style classifier. That means the content encoder learns to actively hide style information, because any style evidence that helps the adversary will hurt the encoder.",
            "I would also stress that architecture is not enough by itself. A paper can claim to have a content code and a style code, but unless it audits style leakage, the split might be much weaker than it looks. That is why later our own Lab 1 does not just train a model; it audits whether the split is actually behaving as intended.",
        ],
        "transition": "Disentanglement is necessary, but it is not sufficient. We still need a meaningful target space to steer the generation.",
    },
    {
        "num": 11,
        "title": "Tool 2: Target Style Spaces",
        "time": "2-3 min",
        "goal": "Explain why the target genre needs to be represented as a structured latent object rather than a flat label.",
        "equations": [
            r"\mathrm{target160}=\mathrm{normalize}([2.0 z_s \Vert 1.0 d_{32}])",
            r"\mu_g=\frac{1}{|S_g|}\sum_{x\in S_g} v(x)",
        ],
        "script": [
            "On this slide I would tell the audience that a target genre should not be treated as a single exemplar or as a one-hot class label. It is better understood as a region in latent space, ideally one that is stable enough to act as a blueprint for generation.",
            "The first line says that we build a style vector by combining a learned style embedding with a smaller set of robust descriptors. The second line says that we can summarize a genre using a centroid over the vectors belonging to that genre. These equations are simple, but the important idea is conceptual: the target should be a geometric object in a meaningful space, not merely a token.",
            "The practical reason is that a generator can only follow a condition if the condition is informative and stable. If the style space is noisy or collapsed, then even a strong generator will seem weak, because the control signal itself is poor.",
        ],
        "transition": "Once we have a target space, the next question is how to inject that target into a model effectively.",
    },
    {
        "num": 12,
        "title": "Tool 3: Conditioning Mechanisms",
        "time": "2 min",
        "goal": "Explain why FiLM and AdaIN are useful and how conditioning depth affects generation quality.",
        "equations": [
            r"\mathrm{FiLM}(F;c)=\gamma(c)\odot F+\beta(c)",
            r"\mathrm{AdaIN}(F;s)=(1+\gamma(s))\odot \mathrm{IN}(F)+\beta(s)",
        ],
        "script": [
            "I would explain FiLM as feature-wise linear modulation. Instead of concatenating a condition once at the input and hoping the network remembers it, FiLM allows the condition to modulate intermediate activations throughout the network. That gives the model more persistent and fine-grained control.",
            "AdaIN is similar in spirit but acts more like a style-control mechanism over normalized features. Conceptually, FiLM is often useful when the condition should guide structure across depth, while AdaIN is often useful when the condition acts more like timbre or texture. The point is not that one is universally better, but that the architecture should place each kind of information where it has the right inductive effect.",
            "This becomes relevant later when we talk about the diffusion branch of DGGR, where we intentionally separate time and content modulation from style modulation.",
        ],
        "transition": "Conditioning only matters if the generator itself is capable of producing high-fidelity audio, so next we compare generator families.",
    },
    {
        "num": 13,
        "title": "Tool 4: Generator Families",
        "time": "2 min",
        "goal": "Compare the main generation families and motivate why the project uses more than one.",
        "script": [
            "For this slide I would quickly compare the major families. Waveform GANs offer parallel generation but can be unstable and often struggle with local coherence. Spectral GANs improve phase handling, but they depend on a good inversion or vocoder path. Neural codec editors are attractive because the codec decoder acts as a realism prior; instead of generating raw waveforms from scratch, we edit latents that already decode into plausible audio. Diffusion models are slower, but they provide strong controllability and often more stable optimization.",
            "The important message is that there is no universally best generator. Different choices optimize different tradeoffs. That is why our project ended up with both a codec branch and a diffusion branch rather than trying to force one model family to do everything.",
        ],
        "transition": "The diffusion family deserves its own slide, because guidance and denoising have become especially important in modern conditional audio generation.",
    },
    {
        "num": 14,
        "title": "Tool 5: Diffusion and Guidance",
        "time": "3 min",
        "goal": "Explain the forward process, reverse process, and classifier-free guidance in a way that connects directly to the project.",
        "equations": [
            r"q(x_t|x_0)=\mathcal{N}(\sqrt{\bar{\alpha}_t}x_0, (1-\bar{\alpha}_t)I)",
            r"\hat{\epsilon}_{cfg}=\epsilon_\theta(x_t,t,\varnothing)+w[\epsilon_\theta(x_t,t,c)-\epsilon_\theta(x_t,t,\varnothing)]",
        ],
        "script": [
            "I would tell the audience that diffusion models work by defining a forward process that gradually corrupts data with noise and then learning a reverse process that reconstructs the data. The first equation describes the noisy state at timestep t as a Gaussian around the original sample with variance controlled by the noise schedule.",
            "The second equation is classifier-free guidance. The model learns both conditional and unconditional behaviors, and at sampling time we interpolate between them with a weight w. That weight becomes a practical knob for how strongly we want to push the sample toward the target condition.",
            "For style transfer, this is useful because edit strength is not binary. If we guide too weakly, the genre shift is barely audible. If we guide too strongly, we can lose content or introduce artifacts. Diffusion is attractive because it gives us an explicit mechanism to manage that balance.",
        ],
        "transition": "That brings us to the last general tool: coherence over longer horizons.",
    },
    {
        "num": 15,
        "title": "Tool 6: Long-Form Coherence",
        "time": "3 min",
        "goal": "Show why long-form generation is a different problem from short-form generation and introduce source anchoring and overlap locking.",
        "equations": [
            r"x_t^{(k)}=\alpha(t)x_0^{(k)}+\sigma(t)\epsilon",
            r"\mathrm{prefix}(x_t^{(k)})\leftarrow \mathrm{tail}(x_t^{(k-1)})",
        ],
        "script": [
            "I would explain that once we move from isolated short clips to full songs, the problem changes qualitatively. A system can sound good for five seconds and still fail badly over a few minutes because local errors accumulate. That is why long-form generation is not just a scaling problem; it is a coherence problem.",
            "The first equation is the anchoring idea. Instead of starting each chunk from pure noise, we start from a noised version of the source chunk. That preserves macro-structure. The second equation is the overlap-locking idea: during sampling, we constrain the prefix of the current chunk to match the tail of the previous one, rather than hoping a later crossfade will hide the mismatch.",
            "I would stress that this is a systems perspective on generation. Long-form quality depends on chunking policy, anchoring, overlap constraints, vocoding, and stabilization together.",
        ],
        "transition": "Before switching into the project section, I would summarize the main design rules that follow from the lecture.",
    },
    {
        "num": 16,
        "title": "Field-Level Design Rules",
        "time": "2 min",
        "goal": "Condense the lecture half into design principles that naturally motivate the DGGR architecture.",
        "script": [
            "This slide is where I would synthesize the general lecture into a few rules. First, measure representation quality separately from generation quality. Second, keep listening in the evaluation loop. Third, prefer staged architectures when the problem is clearly multilevel. And fourth, treat long-form generation as a systems problem rather than as a single-model problem.",
            "I would also say what to avoid: do not trust a high genre score as proof of a musically convincing transfer, do not assume crossfades solve coherence, and do not ignore dataset-source confounds when reading metrics. These are exactly the lessons we tried to carry into our own implementation.",
        ],
        "transition": "With the general framework established, we can now shift into DGGR as a case study.",
    },
    {
        "num": 17,
        "title": "Part II: DGGR as case study",
        "time": "0.5 min",
        "goal": "Signal the transition from literature/tooling overview into the project-specific section.",
        "script": [
            "At this section divider I would explicitly tell the audience that everything from here onward is about our implementation. The question is no longer what the field does in general, but whether our pipeline is a good, defensible realization of those ideas.",
        ],
        "transition": "So the first project-specific question is: what exactly did we try to build?",
    },
    {
        "num": 18,
        "title": "DGGR: Project Thesis",
        "time": "2 min",
        "goal": "State the architecture-level thesis of the project and connect it back to the lecture framework.",
        "script": [
            "Here I would define DGGR as a staged genre-remastering pipeline. The goal is not just to recolor the source audio, but to deconstruct the source into a content representation, define a target genre blueprint, reconstruct realistic audio under that target, and then extend the process to long-form coherence.",
            "I would map the labs directly onto the lecture concepts. Lab 1 corresponds to disentanglement. Lab 2 corresponds to target-space construction. Lab 3 corresponds to reconstruction. Lab 4 corresponds to long-form coherence. Lab 5 corresponds to perceptual validation. Presenting it this way makes the architecture easier to justify because every lab has a conceptual role, not just an implementation role.",
        ],
        "transition": "The next slide shows the full pipeline in one place so the audience can see how the labs connect.",
    },
    {
        "num": 19,
        "title": "DGGR Pipeline Overview",
        "time": "2 min",
        "goal": "Give a single systems-level map of the whole project before diving into each lab.",
        "script": [
            "On this slide I would walk from left to right through the pipeline. The source audio enters Lab 1, where we extract z_content, z_style, and a music-gate decision. Lab 2 converts style-bearing information into a target-space representation and genre centroids. Then Lab 3 branches into two reconstruction paths: a codec-latent translator and a diffusion-based generator. Finally, Lab 4 extends the system to long-form generation using overlap locking and re-anchoring.",
            "The thesis statement I would emphasize is that the project is staged because each lab is solving one identifiable problem from the lecture half. That makes the architecture modular, debuggable, and scientifically easier to interpret.",
        ],
        "transition": "Before we celebrate those modules, though, we need to be honest about the data and labels they depend on.",
    },
    {
        "num": 20,
        "title": "Data Universe and Label Risks",
        "time": "2 min",
        "goal": "Acknowledge dataset and labeling confounds up front so the results are interpreted responsibly.",
        "script": [
            "This is where I would show that our workflow uses multiple kinds of material: baroque or classical renders, hip-hop material, lo-fi material, open-domain CC0 music, and speech negatives for the gate. I would also mention the observed genre buckets in the current calibration run so the audience understands the class distribution we are working with.",
            "Then I would stress the key risk: genre labels can correlate with dataset source, recording conditions, or collection-specific artifacts. That means a model may partially succeed for the wrong reason. For example, it might classify a clip as belonging to a genre because of recording style or dataset fingerprint rather than because it truly learned portable musical style.",
            "I think this honesty is important in a project talk. It raises the technical standard of the presentation and makes later claims more credible, because we are signaling that we understand the limitations of our own evidence.",
        ],
        "transition": "With that caveat on the table, we can look at the first implemented module: the deconstruction encoder.",
    },
    {
        "num": 21,
        "title": "Lab 1: Deconstruction Encoder",
        "time": "2-3 min",
        "goal": "Explain the architecture of Lab 1 and why it is foundational for the rest of the project.",
        "script": [
            "For Lab 1 I would describe the input first: a 96-bin log-mel spectrogram extracted from a fixed five-second chunk. The backbone is a compact Conv2D encoder with stride-based downsampling, global pooling, and a shared projection layer. On top of that backbone, we place multiple heads: z_content, z_style, a style probe, an adversarial style-from-content probe, and the music gate.",
            "The important conceptual point is that Lab 1 is not just feature extraction. It is the place where we try to build a musically useful decomposition of the source. If this lab fails, later labs inherit entangled representations and every downstream result becomes harder to interpret.",
            "I would also mention the training curriculum. Early phases focus more heavily on content separation, and later phases increase adversarial pressure and sharpen the gate. That staged training reflects the fact that disentanglement usually benefits from curriculum rather than from a flat loss schedule.",
        ],
        "transition": "Architecture is one thing, but the more important question is whether the lab actually met its intended scientific thresholds.",
    },
    {
        "num": 22,
        "title": "Lab 1: Measured Outcome",
        "time": "2 min",
        "goal": "Interpret the Lab 1 audit metrics and connect them to the project goals.",
        "script": [
            "Here I would report the three key audit numbers directly: style probe accuracy of 0.9417, content leakage above baseline of 0.1083, and gate ROC AUC of 0.9299. All three clear the thresholds we defined in the project plan.",
            "Then I would interpret them carefully. The style probe tells us that z_style is strongly style informative. The low leakage score suggests that z_content is substantially style suppressed. And the gate ROC AUC tells us that the model can rank music against non-music reliably. I would also mention that threshold calibration still matters. High AUC does not mean we get every precision and recall tradeoff for free.",
            "The broader point is that Lab 1 provides enough evidence to justify using these latents downstream. We are not treating the disentanglement claim as an article of faith; we are treating it as an audited property.",
        ],
        "transition": "Once the source is deconstructed, the next problem is how to define the target genre in a way the generator can actually use.",
    },
    {
        "num": 23,
        "title": "Lab 2: Building the Target160 Style Space",
        "time": "2-3 min",
        "goal": "Explain how the target vector space is constructed and what the visualization is intended to prove.",
        "script": [
            "On this slide I would describe the target160 representation. It combines a 128-dimensional learned style vector with a 32-dimensional descriptor made from summary statistics over mel bands. The reason for adding the descriptor portion is to anchor the learned embedding with simple, robust spectral statistics.",
            "I would explain the t-SNE plot carefully. It is not the final objective, but it is a useful sanity check: if the clusters are obviously collapsed or fully mixed, then the target space is unlikely to be useful for controlled generation. What we want to see is that genres occupy separable regions, even if the separation is not perfect.",
            "I would also remind the audience that this figure comes from our own calibration artifacts, not from a paper. That matters because it shows the slide is connected to the actual repo evidence rather than to a generic literature illustration.",
        ],
        "transition": "The quantitative validation slide then tells us whether that visual separation holds up numerically.",
    },
    {
        "num": 24,
        "title": "Lab 2: Validation",
        "time": "2 min",
        "goal": "Present the target-space metrics and explain why they matter more than they may appear to at first glance.",
        "script": [
            "The main numbers here are linear probe accuracy of 0.8554, nearest-centroid accuracy of 0.8514, and a cosine-silhouette score of 0.4939. The silhouette threshold we set was 0.45, so the current run clears it with margin.",
            "I would explain why this slide matters so much. If the target space is poorly organized, then later generation quality can be limited by a weak control signal rather than by a weak generator. In other words, Lab 2 is upstream of Lab 3. The point is not simply to get a pretty embedding, but to build a geometry that a generator can actually follow.",
            "This is a useful moment in the talk to emphasize scientific discipline: good generation later in the pipeline depends on good representation earlier in the pipeline.",
        ],
        "transition": "Once both source and target representations are in place, we can discuss reconstruction, starting with the strongest current branch: the codec-latent translator.",
    },
    {
        "num": 25,
        "title": "Lab 3A: Codec-Latent Translator",
        "time": "3 min",
        "goal": "Explain the codec branch architecture and the key design decision that improved style shift strength.",
        "script": [
            "I would begin by contrasting this branch with raw waveform generation. Instead of synthesizing raw audio from scratch, we operate on EnCodec quantized embeddings. The translator uses a Conv1D in-projection, a hidden width of 256, and a stack of FiLM-conditioned residual blocks before projecting back to codec-latent space.",
            "The crucial design decision was to allow direct-output mode rather than forcing every output to remain a residual edit of the source latent. Residual mode is safer, but it also acts like an identity leash. It encourages the model to stay too close to the source. Direct-output mode gave the model more freedom to perform a stronger genre remaster.",
            "I would say explicitly that EnCodec helps because its decoder acts as a realism prior. The translator does not need to discover the basic physics of waveform generation from scratch. It only needs to learn how to move through a latent space that already decodes into plausible audio.",
        ],
        "transition": "The next slide explains how this branch was trained and why the best run was significantly better than earlier attempts.",
    },
    {
        "num": 26,
        "title": "Lab 3A: Loss Stack and Best Run",
        "time": "3 min",
        "goal": "Explain the multi-loss training stack and interpret the strongest codec-branch result.",
        "equations": [
            r"\mathcal{L}_{codec}=\lambda_{adv}\mathcal{L}_{GAN}+\lambda_{fm}\mathcal{L}_{feat}+\lambda_1\|q_{hat}-q^*\|_1+\lambda_{mrstft}\mathcal{L}_{MRSTFT}+\lambda_c\mathcal{L}_{content}+\lambda_s\mathcal{L}_{style}+\lambda_p\mathcal{L}_{push}",
        ],
        "script": [
            "I would walk through the loss stack in conceptual blocks rather than in raw notation. The adversarial and feature-matching terms push realism. The L1 and MR-STFT terms keep the reconstruction acoustically grounded. The content and style terms enforce the remastering objective itself. The push term encourages movement away from trivial identity solutions.",
            "Then I would explain why run1055 matters. It achieved an MPS of 0.9565, style confidence of 0.8940, and style accuracy of 0.9492. In other words, it preserved melodic content very strongly while still producing a target-style shift that exceeded our intended threshold.",
            "My interpretation would be that three things mattered together: MERT-based conditioning improved the target-style signal, direct-output mode increased edit freedom, and explicit content preservation losses prevented that stronger edit from collapsing melody. This branch is currently the strongest short-form evidence path in the repo.",
        ],
        "transition": "The diffusion branch is more ambitious and more difficult, so the next slides explain both its architecture and its current limitations.",
    },
    {
        "num": 27,
        "title": "Lab 3B: Diffusion V2",
        "time": "3 min",
        "goal": "Explain the diffusion architecture in a project-specific way and connect it to the earlier lecture slide on conditioning.",
        "script": [
            "On this slide I would first decode the input tensor for the audience. The model does not consume only a noisy mel spectrogram. It also receives chroma channels, an onset channel, and a beat-grid channel. So even before generation, we are already telling the model something about harmonic and rhythmic structure.",
            "Then I would describe the backbone: a UNet with channel progression 64, 128, 256, 256, two residual blocks per level, low-resolution attention, and an EMA shadow model for more stable sampling. The important architectural idea, however, is the condition split. Time and content are injected through FiLM, while style is injected through dedicated StyleAdaIN blocks.",
            "I would connect this back to the lecture directly. In the lecture we said that style and structure often benefit from different modulation paths. This branch applies that design principle explicitly. It is trying to stop timbral control from overwhelming melodic control.",
        ],
        "transition": "The next slide shows a very important empirical lesson from this branch: numeric optimization and perceptual quality do not peak at the same checkpoint.",
    },
    {
        "num": 28,
        "title": "Diffusion Training Behavior",
        "time": "2 min",
        "goal": "Explain why epoch selection in generative audio cannot rely on validation loss alone.",
        "script": [
            "I would tell the audience that the validation-loss curve kept improving through later epochs, but the best-sounding checkpoint in practice was epoch 6. This is exactly the kind of result that makes generative audio tricky to evaluate. A smoother, more averaged model can improve scalar loss while losing the sharper stylistic character we actually want to hear.",
            "So the takeaway is methodological. In this domain, listening checkpoints are not just demo material. They are part of the scientific loop. If you ignore perceptual evaluation until the very end, you can select a model that is mathematically cleaner but musically less convincing.",
            "I would also be honest that this branch is currently less reliable than the codec branch in our repo. It remains interesting because it offers more generative freedom and is the branch that supports our long-form experiments.",
        ],
        "transition": "That naturally leads into Lab 4, where the problem is no longer only realism but coherence across an entire song.",
    },
    {
        "num": 29,
        "title": "Lab 4: Long-Form Coherence Diagnostics",
        "time": "3 min",
        "goal": "Explain the long-form setup, the diagnostics, and what the current metrics do and do not prove.",
        "script": [
            "For this slide I would describe the full-song test as a stress test of the system rather than as a final solved result. The setup used 64 overlapping chunks over a 160-second track. The purpose was to see whether overlap locking and source anchoring gave us measurable continuity across the entire generation process.",
            "The two key diagnostics are boundary mel MSE and boundary discontinuity in decibels. Lower values indicate better continuity across chunk boundaries. Our current means, roughly 0.0018 for boundary mel MSE and 2.87 decibels for discontinuity, tell us that the seam problem is at least being managed quantitatively.",
            "But I would be careful not to oversell that result. These metrics do not prove that the long-form output sounds fully convincing. What they really tell us is that obvious hard seams are not the dominant remaining issue. The harder issue is slow accumulation of warble and static as style edits compound across time. That is why I call long-form generation a systems problem.",
        ],
        "transition": "Before moving to audio demos, it helps to summarize the metrics across the whole project in one place.",
    },
    {
        "num": 30,
        "title": "Quantitative Summary Across Labs",
        "time": "2 min",
        "goal": "Synthesize the project results and make clear which claims are strong and which are still provisional.",
        "script": [
            "Here I would use the dashboard to make a disciplined summary. Lab 1 and Lab 2 both clear their planned thresholds. That means the representational side of the project is in good shape. The codec branch also clears its short-form gates with strong melody preservation and strong style metrics. Those are the strongest claims we can make confidently.",
            "Then I would explicitly separate those strong claims from the open ones. The diffusion branch is promising but less mature. The long-form branch shows measurable coherence progress but still suffers from perceptual artifact accumulation. And Lab 5, the human listening study, is the major missing piece if we want to move from technical plausibility to stronger empirical validation.",
        ],
        "transition": "Once the audience has that quantitative summary in mind, we can listen to examples in a more structured way.",
    },
    {
        "num": 31,
        "title": "Demo Slide: Short-Form Codec Examples",
        "time": "3-4 min",
        "goal": "Guide the audience on how to listen to the codec examples rather than playing them without interpretation.",
        "audio_cues": [
            "Play `codec_src1_tgt3.mp3` first: CC0 other to lo-fi.",
            "Play `codec_src2_tgt1.mp3` second: hip-hop to open-domain other.",
            "Play `codec_src3_tgt0.mp3` third: lo-fi to baroque.",
        ],
        "script": [
            "I would not just play the clips silently. I would first tell the audience what to listen for. The first question is whether the melody or core phrase identity survives. The second is whether the target genre is actually perceptible. The third is whether the audio sounds plausible in its own right rather than obviously synthetic or unstable.",
            "Then I would introduce each clip briefly. For example, CC0 other to lo-fi, hip-hop to open-domain other, and lo-fi to baroque. The point of using several directions is to show that the branch is not limited to one especially easy source-target pair.",
            "After each clip, I would make one concise interpretive remark. For example: this one preserves the phrase contour well but makes only a moderate timbral shift, or this one makes a strong shift but introduces some synthetic texture on attacks. Those remarks help the audience compare strengths and weaknesses actively instead of passively listening.",
        ],
        "transition": "The next demo slide changes the listening criteria, because diffusion and long-form audio raise different questions than short-form codec remastering.",
    },
    {
        "num": 32,
        "title": "Demo Slide: Diffusion and Long-Form Excerpts",
        "time": "3-4 min",
        "goal": "Frame the listening task for the diffusion and long-form clips so the audience hears the right phenomena.",
        "audio_cues": [
            "Play `diffusion_v2_epoch6_gen0.mp3` first as the short diffusion example.",
            "Play `longform_source_excerpt.mp3` second as the source reference for the 30s to 50s segment.",
            "Play `longform_remaster_excerpt.mp3` third as the matching long-form remaster excerpt.",
        ],
        "script": [
            "Here I would explicitly change the listening instructions. For the diffusion clip, the key question is how much perceptual freedom the model seems to have and whether the result sounds expressive rather than oversmoothed. For the long-form pair, the key questions are continuity of phrasing, seam quality, and artifact buildup over time.",
            "I would first play the short diffusion example as a demonstration of the branch's generative ambition. Then I would play the long-form source excerpt and the corresponding remastered excerpt back to back. The audience should not be expected to judge this clip only on style accuracy; they should also listen for continuity and stability across the whole segment.",
            "If the clip shows mild boundary continuity but growing timbral instability, I would say that directly. The point of the demo is not to hide the limitations. It is to show that the system is capable of meaningful long-form remastering behavior while still having room for refinement.",
        ],
        "transition": "After the demos, the last two slides should tell the audience what is genuinely novel and what still needs to be finished.",
    },
    {
        "num": 33,
        "title": "Creative Elements and Limitations",
        "time": "2 min",
        "goal": "Articulate the project's novelty in a defensible way while staying honest about its present weaknesses.",
        "script": [
            "For the creative-elements portion, I would avoid claiming that any single block is completely unprecedented. Instead, I would argue that the novelty lies in the combination and the staging. The project frames genre remastering as a structure-first problem, implements an explicit deconstruction-calibration-reconstruction-coherence pipeline, and makes concrete engineering decisions such as direct-output codec translation, MERT-based conditioning, and prefix-lock sampling constraints.",
            "That is important in the context of the course rubric, because creative elements do not have to mean inventing an entirely new model family from scratch. They can also mean building a more effective, more interpretable, and more technically justified system out of known components in a way that addresses real bottlenecks.",
            "On the limitations side, I would be direct: dataset-source leakage remains a risk, diffusion still trails the codec branch in short-form reliability, and long-form outputs still accumulate warble or static under stronger edits. I think stating those limitations clearly actually strengthens the presentation.",
        ],
        "transition": "The final slide then turns those limitations into a concrete next step rather than leaving the talk in an open-ended state.",
    },
    {
        "num": 34,
        "title": "Lab 5 and Conclusion",
        "time": "2 min",
        "goal": "End the talk by reframing the remaining work as a clear, testable human-evaluation question.",
        "script": [
            "I would open the final slide by saying that the most important remaining question is no longer whether the system can run, or even whether it can produce promising quantitative results. The remaining question is whether listeners actually hear what we intend them to hear. That is the purpose of Lab 5.",
            "So I would phrase the evaluation goals explicitly. Can listeners identify the target genre? Can they still recognize the source melody? Do they prefer our architecture-driven remaster to a simpler filter-style baseline? Those questions turn the remaining work into a well-defined experiment rather than a vague promise of future improvement.",
            "To close, I would summarize the entire talk in one sentence: we started from the general problem of meaningful genre transfer, built a staged repo-backed system that already clears several key technical gates, and now need formal perceptual validation to determine how convincing the full remastering experience is to human listeners.",
        ],
        "transition": "At that point I would stop and invite questions, ideally starting from architecture tradeoffs, evaluation design, or where the long-term path should go next.",
    },
]


def esc(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    text = text.replace("≥", r"$\geq$")
    text = text.replace("→", r"$\rightarrow$")
    return text


def directify(text: str) -> str:
    replacements = [
        (r"^I would open by telling the audience that ", ""),
        (r"^I would also say, right away, that ", ""),
        (r"^On this slide I would tell the audience that ", ""),
        (r"^Then I would preview ", "Then preview "),
        (r"^At this section divider I would pause and explicitly say that ", ""),
        (r"^At this section divider I would explicitly tell the audience that ", ""),
        (r"^Here I would start with ", "Start with "),
        (r"^Then I would emphasize that ", ""),
        (r"^Finally, I would list ", "List "),
        (r"^For this slide I would define content as ", "Content consists of "),
        (r"^I would define style as ", "Style consists of "),
        (r"^I would read the first expression very simply: ", ""),
        (r"^I would explain that ", ""),
        (r"^On this slide I would summarize the field in phases\. ", "The field can be summarized in phases. "),
        (r"^I would use this slide to make a simple but important point: ", ""),
        (r"^I would explain the notation first\. ", "Start with the notation. "),
        (r"^I would also stress that ", ""),
        (r"^I would explain FiLM as ", "FiLM is "),
        (r"^For this slide I would quickly compare ", "Compare "),
        (r"^I would tell the audience that ", ""),
        (r"^I would stress that ", ""),
        (r"^This slide is where I would synthesize ", "This slide synthesizes "),
        (r"^I would also say what to avoid: ", "Avoid the following mistakes: "),
        (r"^Here I would define DGGR as ", "DGGR is "),
        (r"^I would map the labs directly onto the lecture concepts\. ", "Map the labs directly onto the lecture concepts. "),
        (r"^On this slide I would walk from left to right through the pipeline\. ", "Walk from left to right through the pipeline. "),
        (r"^The thesis statement I would emphasize is that ", "The key thesis is that "),
        (r"^This is where I would show that ", ""),
        (r"^Then I would stress the key risk: ", "The key risk is the following: "),
        (r"^For Lab 1 I would describe the input first: ", "Start with the input: "),
        (r"^I would also mention the training curriculum\. ", "Mention the training curriculum as well. "),
        (r"^Here I would report ", "Report "),
        (r"^Then I would interpret them carefully\. ", "Interpret those numbers carefully. "),
        (r"^On this slide I would describe ", "This slide describes "),
        (r"^I would explain the t-SNE plot carefully\. ", "Interpret the t-SNE plot carefully. "),
        (r"^I would also remind the audience that ", ""),
        (r"^I would explain why this slide matters so much\. ", "This slide matters because "),
        (r"^I would begin by contrasting this branch with raw waveform generation\. ", "Begin by contrasting this branch with raw waveform generation. "),
        (r"^I would say explicitly that ", ""),
        (r"^I would walk through the loss stack in conceptual blocks rather than in raw notation\. ", "Walk through the loss stack in conceptual blocks rather than raw notation. "),
        (r"^Then I would explain why ", "Then explain why "),
        (r"^On this slide I would first decode the input tensor for the audience\. ", "Start by decoding the input tensor. "),
        (r"^Then I would describe the backbone: ", "Then describe the backbone: "),
        (r"^I would connect this back to the lecture directly\. ", "Connect this back to the lecture directly. "),
        (r"^I would also be honest that ", ""),
        (r"^For this slide I would describe ", "This slide describes "),
        (r"^But I would be careful not to oversell that result\. ", "Do not oversell that result. "),
        (r"^Here I would use the dashboard to make a disciplined summary\. ", "Use the dashboard to make a disciplined summary. "),
        (r"^Then I would explicitly separate ", "Then separate "),
        (r"^I would not just play the clips silently\. ", "Do not just play the clips silently. "),
        (r"^I would first tell the audience what to listen for\. ", "Start by telling the audience what to listen for. "),
        (r"^Then I would introduce each clip briefly\. ", "Then introduce each clip briefly. "),
        (r"^After each clip, I would make one concise interpretive remark\. ", "After each clip, make one concise interpretive remark. "),
        (r"^Here I would explicitly change the listening instructions\. ", "Change the listening instructions explicitly. "),
        (r"^I would first play ", "First play "),
        (r"^If the clip shows mild boundary continuity but growing timbral instability, I would say that directly\. ", "If the clip shows mild boundary continuity but growing timbral instability, say that directly. "),
        (r"^For the creative-elements portion, I would avoid claiming that ", "For the creative-elements portion, avoid claiming that "),
        (r"^On the limitations side, I would be direct: ", "On the limitations side, be direct: "),
        (r"^I would open the final slide by saying that ", ""),
        (r"^So I would phrase the evaluation goals explicitly\. ", "Phrase the evaluation goals explicitly. "),
        (r"^To close, I would summarize ", "To close, summarize "),
        (r"^At that point I would stop and invite questions", "At that point, invite questions"),
        (r"^Before switching into the project section, I would summarize ", "Before switching into the project section, summarize "),
    ]
    out = text.strip()
    for pattern, repl in replacements:
        out = re.sub(pattern, repl, out)
    out = re.sub(r"^I think this honesty is important in a project talk\. ", "This honesty is important in a project talk. ", out)
    out = re.sub(r"\bI think stating those limitations clearly actually strengthens the presentation\.", "Stating those limitations clearly actually strengthens the presentation.", out)
    out = re.sub(r"\bI call long-form generation a systems problem\.", "Long-form generation is best understood as a systems problem.", out)
    out = re.sub(r"\bI would\b", "", out)
    out = re.sub(r"\s+", " ", out).strip()
    if out:
        out = out[0].upper() + out[1:]
    return out


def equation_explanations_for(slide_num: int) -> list[str]:
    return EQUATION_EXPLANATIONS.get(slide_num, [])


def build_tex() -> None:
    lines: list[str] = [
        r"\documentclass[11pt]{article}",
        r"\pdfminorversion=4",
        r"\pdfobjcompresslevel=0",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{lmodern}",
        r"\usepackage{microtype}",
        r"\usepackage{amsmath}",
        r"\usepackage{amssymb}",
        r"\usepackage{booktabs}",
        r"\usepackage{xcolor}",
        r"\definecolor{Ink}{HTML}{1F2937}",
        r"\definecolor{Muted}{HTML}{637381}",
        r"\title{DGGR Final Presentation Speaker Script}",
        r"\author{Sahara Kaul \and Kelsey Pattison \and Ahmed Sajid}",
        r"\date{CMPUT 414, Winter 2026}",
        r"\begin{document}",
        r"\maketitle",
        r"\tableofcontents",
        r"\newpage",
        r"\section*{How To Use This Script}",
        r"This document is a detailed speaker script keyed to the 34-slide deck \texttt{dggr\_lecture\_deck.pptx}.",
        r"It is intentionally more detailed than what should be spoken verbatim. The intended use is as a flow document:",
        r"\begin{itemize}",
        r"\item read each slide's \textbf{goal} before practicing;",
        r"\item use the script paragraphs as phrasing material rather than as a mandatory word-for-word recitation;",
        r"\item keep the \textbf{transition} sentence in mind so the talk moves smoothly slide to slide.",
        r"\end{itemize}",
        r"\newpage",
    ]
    current_section = None
    for slide in SLIDES:
        if slide["num"] == 1:
            current_section = "Opening"
            lines.append(r"\newpage")
            lines.append(r"\section{Opening}")
        elif slide["num"] == 3:
            current_section = "Part I: General Lecture"
            lines.append(r"\newpage")
            lines.append(r"\section{Part I: General Lecture}")
        elif slide["num"] == 17:
            current_section = "Part II: DGGR Case Study"
            lines.append(r"\newpage")
            lines.append(r"\section{Part II: DGGR Case Study}")

        lines.append(f"\\subsection*{{Slide {slide['num']}: {esc(slide['title'])}}}")
        lines.append(f"\\textbf{{Goal:}} {esc(slide['goal'])}")
        if slide.get("equations"):
            lines.append(r"\paragraph{Equations to mention}")
            lines.append(r"\begin{align*}")
            for idx, eq in enumerate(slide["equations"]):
                suffix = r"\\" if idx < len(slide["equations"]) - 1 else ""
                lines.append(eq + suffix)
            lines.append(r"\end{align*}")
            explanations = equation_explanations_for(slide["num"])
            if explanations:
                lines.append(r"\paragraph{Equation explanation}")
                lines.append(r"\begin{itemize}")
                for item in explanations:
                    lines.append(r"\item " + esc(item))
                lines.append(r"\end{itemize}")
        if slide.get("audio_cues"):
            lines.append(r"\paragraph{Audio cues}")
            lines.append(r"\begin{itemize}")
            for cue in slide["audio_cues"]:
                lines.append(r"\item " + esc(cue))
            lines.append(r"\end{itemize}")
        lines.append(r"\paragraph{Talk track}")
        for para in slide["script"]:
            lines.append(esc(directify(para)))
            lines.append("")
        lines.append(r"\paragraph{Transition}")
        lines.append(esc(directify(slide["transition"])))
        lines.append("")
    lines.append(r"\end{document}")
    TEX_OUT.write_text("\n".join(lines), encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.add_run("DGGR Final Presentation Speaker Script")
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("CMPUT 414, Winter 2026").italic = True

    doc.add_heading("How To Use This Script", level=1)
    doc.add_paragraph(
        "This document is keyed to the 34-slide deck dggr_lecture_deck.pptx. "
        "It is intentionally more detailed than what should be spoken verbatim."
    )
    for bullet in [
        "Read each slide goal before practicing.",
        "Use the main script as phrasing material rather than as a strict word-for-word script.",
        "Use the transition line to keep the flow between slides smooth.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    for slide in SLIDES:
        if slide["num"] == 1:
            doc.add_page_break()
            doc.add_heading("Opening", level=1)
        elif slide["num"] == 3:
            doc.add_page_break()
            doc.add_heading("Part I: General Lecture", level=1)
        elif slide["num"] == 17:
            doc.add_page_break()
            doc.add_heading("Part II: DGGR Case Study", level=1)

        doc.add_heading(f"Slide {slide['num']}: {slide['title']}", level=2)
        p = doc.add_paragraph()
        p.add_run("Goal: ").bold = True
        p.add_run(slide["goal"])
        if slide.get("equations"):
            p = doc.add_paragraph()
            p.add_run("Equations to mention:").bold = True
            for eq in slide["equations"]:
                doc.add_paragraph(eq, style="List Bullet")
            explanations = equation_explanations_for(slide["num"])
            if explanations:
                p = doc.add_paragraph()
                p.add_run("Equation explanation:").bold = True
                for item in explanations:
                    doc.add_paragraph(item, style="List Bullet")
        if slide.get("audio_cues"):
            p = doc.add_paragraph()
            p.add_run("Audio cues:").bold = True
            for cue in slide["audio_cues"]:
                doc.add_paragraph(cue, style="List Bullet")
        p = doc.add_paragraph()
        p.add_run("Talk track:").bold = True
        for para in slide["script"]:
            doc.add_paragraph(directify(para))
        p = doc.add_paragraph()
        p.add_run("Transition: ").bold = True
        p.add_run(directify(slide["transition"]))

    doc.save(DOCX_OUT)


def main() -> None:
    build_tex()
    build_docx()
    print(TEX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
